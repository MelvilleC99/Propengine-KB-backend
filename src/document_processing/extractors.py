"""Document Extractors - Extract text and structure from uploaded documents"""

import io
import logging
import re
from typing import Dict, Any, List, Optional
from dataclasses import dataclass

logger = logging.getLogger(__name__)


@dataclass
class ExtractedSection:
    """Represents a section extracted from a document"""
    heading: str
    content: str
    level: int  # Heading level (1 = H1, 2 = H2, etc.)
    page_start: Optional[int] = None


@dataclass
class ExtractionResult:
    """Result of document extraction"""
    full_text: str
    sections: List[ExtractedSection]
    metadata: Dict[str, Any]
    success: bool
    error: Optional[str] = None


class DocxExtractor:
    """
    Extracts text and structure from DOCX files.
    Preserves heading hierarchy and extracts table content.
    """
    
    def __init__(self):
        try:
            from docx import Document
            self._docx_available = True
        except ImportError:
            logger.warning("python-docx not installed. DOCX extraction will be limited.")
            self._docx_available = False
    
    async def extract(self, file_bytes: bytes, filename: str) -> ExtractionResult:
        """
        Extract text and structure from DOCX file bytes.
        
        Args:
            file_bytes: Raw bytes of the DOCX file
            filename: Original filename for metadata
            
        Returns:
            ExtractionResult with text, sections, and metadata
        """
        try:
            if not self._docx_available:
                return await self._fallback_extract(file_bytes, filename)
            
            return await self._extract_with_structure(file_bytes, filename)
            
        except Exception as e:
            logger.error(f"❌ Failed to extract DOCX: {e}", exc_info=True)
            return ExtractionResult(
                full_text="",
                sections=[],
                metadata={"filename": filename},
                success=False,
                error=str(e)
            )
    
    async def _extract_with_structure(self, file_bytes: bytes, filename: str) -> ExtractionResult:
        """Extract with full structure preservation using python-docx"""
        from docx import Document
        from docx.document import Document as DocType
        from docx.oxml.table import CT_Tbl
        from docx.oxml.text.paragraph import CT_P
        from docx.table import Table
        from docx.text.paragraph import Paragraph
        
        # Load document from bytes
        doc = Document(io.BytesIO(file_bytes))
        
        sections: List[ExtractedSection] = []
        full_text_parts: List[str] = []
        
        current_section_heading = "Introduction"
        current_section_content: List[str] = []
        current_section_level = 1
        
        def process_paragraph(para: Paragraph):
            """Process a paragraph element"""
            nonlocal current_section_heading, current_section_content, current_section_level
            
            text = para.text.strip()
            if not text:
                return
            
            # Check if this paragraph is a heading
            heading_level = self._get_heading_level(para)
            
            if heading_level:
                # Save previous section if it has content
                if current_section_content:
                    sections.append(ExtractedSection(
                        heading=current_section_heading,
                        content="\n".join(current_section_content),
                        level=current_section_level
                    ))
                
                # Start new section
                current_section_heading = text
                current_section_content = []
                current_section_level = heading_level
            else:
                # Skip Word error messages
                if not self._is_word_error_text(text):
                    current_section_content.append(text)
            
            if not self._is_word_error_text(text):
                full_text_parts.append(text)
        
        def process_table(table: Table):
            """Process a table element - extract all cell text"""
            nonlocal current_section_content
            
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text and not self._is_word_error_text(cell_text):
                        # Clean up the text (remove excessive whitespace)
                        cell_text = re.sub(r'\s+', ' ', cell_text)
                        row_texts.append(cell_text)
                
                if row_texts:
                    combined_text = " ".join(row_texts)
                    current_section_content.append(combined_text)
                    full_text_parts.append(combined_text)
        
        # Iterate through document body in order (paragraphs AND tables)
        # This preserves the document order
        for element in doc.element.body:
            if isinstance(element, CT_P):
                para = Paragraph(element, doc)
                process_paragraph(para)
            elif isinstance(element, CT_Tbl):
                table = Table(element, doc)
                process_table(table)
        
        # Don't forget the last section
        if current_section_content:
            sections.append(ExtractedSection(
                heading=current_section_heading,
                content="\n".join(current_section_content),
                level=current_section_level
            ))
        
        # If no sections were detected, create one from all content
        if not sections and full_text_parts:
            sections.append(ExtractedSection(
                heading="Document Content",
                content="\n".join(full_text_parts),
                level=1
            ))
        
        full_text = "\n\n".join(full_text_parts)
        
        # Extract metadata
        metadata = self._extract_metadata(doc, filename, full_text)
        
        logger.info(f"✅ Extracted DOCX: {len(sections)} sections, {len(full_text)} chars, {len(doc.tables)} tables")
        
        return ExtractionResult(
            full_text=full_text,
            sections=sections,
            metadata=metadata,
            success=True
        )
    
    def _is_word_error_text(self, text: str) -> bool:
        """Check if text is a Word error/placeholder/comment message"""
        error_patterns = [
            "no table of contents entries found",
            "error! no table of figures entries found",
            "error! bookmark not defined",
            "error! reference source not found",
            "commented [",  # Word comments like "Commented [SN1]:"
            "comment:",     # Alternative comment format
        ]
        text_lower = text.lower().strip()
        
        # Check for @ mentions (like @username in comments)
        if text_lower.startswith("@"):
            return True
            
        return any(pattern in text_lower for pattern in error_patterns)
    
    def _get_heading_level(self, paragraph) -> Optional[int]:
        """
        Determine if paragraph is a heading and return its level.
        
        Returns:
            Heading level (1-9) or None if not a heading
        """
        style_name = paragraph.style.name.lower() if paragraph.style else ""
        
        # Check for built-in heading styles
        if style_name.startswith("heading"):
            try:
                level = int(style_name.replace("heading", "").strip())
                return level
            except ValueError:
                return 1
        
        # Check for title style
        if style_name == "title":
            return 1
        
        # Check for subtitle
        if style_name == "subtitle":
            return 2
        
        return None
    
    def _extract_metadata(self, doc, filename: str, full_text: str) -> Dict[str, Any]:
        """Extract metadata from document"""
        core_props = doc.core_properties
        
        word_count = len(full_text.split())
        estimated_pages = max(1, word_count // 300)
        
        return {
            "filename": filename,
            "title": core_props.title or filename.rsplit(".", 1)[0],
            "author": core_props.author or "Unknown",
            "created": str(core_props.created) if core_props.created else None,
            "modified": str(core_props.modified) if core_props.modified else None,
            "word_count": word_count,
            "char_count": len(full_text),
            "estimated_pages": estimated_pages,
            "table_count": len(doc.tables)
        }
    
    async def _fallback_extract(self, file_bytes: bytes, filename: str) -> ExtractionResult:
        """Fallback extraction using docx2txt if python-docx not available"""
        try:
            import docx2txt
            
            text = docx2txt.process(io.BytesIO(file_bytes))
            
            sections = [ExtractedSection(
                heading="Document Content",
                content=text,
                level=1
            )]
            
            return ExtractionResult(
                full_text=text,
                sections=sections,
                metadata={
                    "filename": filename,
                    "word_count": len(text.split()),
                    "char_count": len(text),
                    "extraction_method": "fallback"
                },
                success=True
            )
            
        except Exception as e:
            logger.error(f"❌ Fallback extraction failed: {e}")
            return ExtractionResult(
                full_text="",
                sections=[],
                metadata={"filename": filename},
                success=False,
                error=f"Extraction failed: {str(e)}"
            )


class PdfExtractor:
    """
    Extracts text and structure from PDF files.
    Uses PyMuPDF (fitz) for reliable extraction.
    """
    
    def __init__(self):
        try:
            import fitz  # PyMuPDF
            self._fitz_available = True
        except ImportError:
            logger.warning("PyMuPDF not installed. PDF extraction will be limited.")
            self._fitz_available = False
    
    async def extract(self, file_bytes: bytes, filename: str) -> ExtractionResult:
        """
        Extract text and structure from PDF file bytes.
        
        Args:
            file_bytes: Raw bytes of the PDF file
            filename: Original filename for metadata
            
        Returns:
            ExtractionResult with text, sections, and metadata
        """
        try:
            if not self._fitz_available:
                return await self._fallback_extract(file_bytes, filename)
            
            return await self._extract_with_fitz(file_bytes, filename)
            
        except Exception as e:
            logger.error(f"❌ Failed to extract PDF: {e}", exc_info=True)
            return ExtractionResult(
                full_text="",
                sections=[],
                metadata={"filename": filename},
                success=False,
                error=str(e)
            )
    
    async def _extract_with_fitz(self, file_bytes: bytes, filename: str) -> ExtractionResult:
        """Extract PDF content using PyMuPDF"""
        import fitz
        
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        
        sections: List[ExtractedSection] = []
        full_text_parts: List[str] = []
        
        current_section_heading = "Introduction"
        current_section_content: List[str] = []
        current_section_level = 1
        
        for page_num, page in enumerate(doc):
            # Extract text blocks with position info
            blocks = page.get_text("dict")["blocks"]
            
            for block in blocks:
                if block["type"] == 0:  # Text block
                    for line in block.get("lines", []):
                        line_text = ""
                        max_font_size = 0
                        is_bold = False
                        
                        for span in line.get("spans", []):
                            text = span.get("text", "").strip()
                            if text:
                                line_text += text + " "
                                font_size = span.get("size", 12)
                                max_font_size = max(max_font_size, font_size)
                                # Check if bold
                                font_name = span.get("font", "").lower()
                                if "bold" in font_name:
                                    is_bold = True
                        
                        line_text = line_text.strip()
                        if not line_text:
                            continue
                        
                        # Detect headings based on font size and style
                        heading_level = self._detect_heading(line_text, max_font_size, is_bold)
                        
                        if heading_level:
                            # Save previous section
                            if current_section_content:
                                sections.append(ExtractedSection(
                                    heading=current_section_heading,
                                    content="\n".join(current_section_content),
                                    level=current_section_level,
                                    page_start=page_num + 1
                                ))
                            
                            current_section_heading = line_text
                            current_section_content = []
                            current_section_level = heading_level
                        else:
                            current_section_content.append(line_text)
                        
                        full_text_parts.append(line_text)
        
        # Don't forget last section
        if current_section_content:
            sections.append(ExtractedSection(
                heading=current_section_heading,
                content="\n".join(current_section_content),
                level=current_section_level
            ))
        
        # If no sections, create one from all content
        if not sections and full_text_parts:
            sections.append(ExtractedSection(
                heading="Document Content",
                content="\n".join(full_text_parts),
                level=1
            ))
        
        full_text = "\n\n".join(full_text_parts)
        
        metadata = {
            "filename": filename,
            "title": doc.metadata.get("title") or filename.rsplit(".", 1)[0],
            "author": doc.metadata.get("author") or "Unknown",
            "created": doc.metadata.get("creationDate"),
            "modified": doc.metadata.get("modDate"),
            "word_count": len(full_text.split()),
            "char_count": len(full_text),
            "page_count": len(doc),
            "estimated_pages": len(doc)
        }
        
        doc.close()
        
        logger.info(f"✅ Extracted PDF: {len(sections)} sections, {len(full_text)} chars, {metadata['page_count']} pages")
        
        return ExtractionResult(
            full_text=full_text,
            sections=sections,
            metadata=metadata,
            success=True
        )
    
    def _detect_heading(self, text: str, font_size: float, is_bold: bool) -> Optional[int]:
        """
        Detect if text is a heading based on font size and formatting.
        
        Returns heading level (1-3) or None if not a heading.
        """
        # Common heading patterns
        heading_pattern = re.match(r'^(\d+\.)+\s*', text)  # 1.1. or 1.2.3.
        
        # Large font (>14pt) or numbered heading pattern with bold
        if font_size >= 18:
            return 1
        elif font_size >= 14 or (heading_pattern and is_bold):
            return 2
        elif is_bold and len(text) < 100 and not text.endswith('.'):
            # Short bold text that doesn't end with period = likely heading
            return 3
        
        return None
    
    async def _fallback_extract(self, file_bytes: bytes, filename: str) -> ExtractionResult:
        """Fallback extraction using pdfplumber or pypdf"""
        try:
            # Try pdfplumber first
            import pdfplumber
            
            full_text_parts = []
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        full_text_parts.append(text)
            
            full_text = "\n\n".join(full_text_parts)
            
            sections = [ExtractedSection(
                heading="Document Content",
                content=full_text,
                level=1
            )]
            
            return ExtractionResult(
                full_text=full_text,
                sections=sections,
                metadata={
                    "filename": filename,
                    "word_count": len(full_text.split()),
                    "char_count": len(full_text),
                    "extraction_method": "pdfplumber"
                },
                success=True
            )
            
        except ImportError:
            pass
        
        # Try pypdf as last resort
        try:
            from pypdf import PdfReader
            
            reader = PdfReader(io.BytesIO(file_bytes))
            full_text_parts = []
            
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    full_text_parts.append(text)
            
            full_text = "\n\n".join(full_text_parts)
            
            sections = [ExtractedSection(
                heading="Document Content",
                content=full_text,
                level=1
            )]
            
            return ExtractionResult(
                full_text=full_text,
                sections=sections,
                metadata={
                    "filename": filename,
                    "word_count": len(full_text.split()),
                    "char_count": len(full_text),
                    "extraction_method": "pypdf"
                },
                success=True
            )
            
        except Exception as e:
            logger.error(f"❌ PDF fallback extraction failed: {e}")
            return ExtractionResult(
                full_text="",
                sections=[],
                metadata={"filename": filename},
                success=False,
                error=f"PDF extraction failed: {str(e)}"
            )


def get_extractor(filename: str):
    """
    Factory function to get the appropriate extractor based on file extension.
    
    Args:
        filename: Name of the file to extract
        
    Returns:
        Appropriate extractor instance
    """
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    
    if ext == "docx":
        return DocxExtractor()
    elif ext == "pdf":
        return PdfExtractor()
    else:
        raise ValueError(f"Unsupported file type: {ext}. Supported: docx, pdf")
